"""
DOCX text extraction module using python-docx.
"""

import os
from typing import Dict, List, Optional

import docx


class DocxExtractor:
    """
    Extract text content from DOCX files.
    """

    def __init__(self):
        """Initialize the DOCX extractor."""
        pass

    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a DOCX file.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Extracted text content.

        Raises:
            FileNotFoundError: If the file does not exist.
            ValueError: If the file is not a valid DOCX.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            doc = docx.Document(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                full_text.append(para.text)
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            return "\n".join(full_text)
        except Exception as e:
            raise ValueError(f"Error reading DOCX: {e}")
    
    def extract_metadata(self, file_path: str) -> Dict:
        """
        Extract metadata from a DOCX file.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Dictionary containing DOCX metadata.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        try:
            doc = docx.Document(file_path)
            
            # Extract core properties
            core_props = {}
            if doc.core_properties:
                props = doc.core_properties
                core_props = {
                    "author": props.author,
                    "category": props.category,
                    "comments": props.comments,
                    "content_status": props.content_status,
                    "created": props.created,
                    "identifier": props.identifier,
                    "keywords": props.keywords,
                    "language": props.language,
                    "last_modified_by": props.last_modified_by,
                    "last_printed": props.last_printed,
                    "modified": props.modified,
                    "revision": props.revision,
                    "subject": props.subject,
                    "title": props.title,
                    "version": props.version
                }
                
                # Filter out None values
                core_props = {k: v for k, v in core_props.items() if v is not None}
            
            # Add document statistics
            doc_stats = {
                "paragraphs": len(doc.paragraphs),
                "sections": len(doc.sections),
                "tables": len(doc.tables)
            }
            
            return {**core_props, **doc_stats}
        except Exception as e:
            raise ValueError(f"Error extracting DOCX metadata: {e}") 